"""Pattadar Assistant — FastAPI service with SSE streaming chat.

Ported from rhub api/assistant/main.py. Differences:
- Identity: trusts the gateway-injected `x-user-id` header (this service is
  NEVER internet-exposed directly; the gateway strips any spoofed header and
  injects the validated one — same invariant as the api service).
- No Jira auto-attach / Builder (/vkx) paths — rhub-only, deleted.
- Creates its own tables (r_conversations / r_attachments) if absent so a
  fresh database works without a separate schema step.
"""

import json
import logging
import re
from contextlib import asynccontextmanager
from typing import AsyncGenerator
from uuid import UUID

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s %(message)s",
)

import psycopg
from fastapi import FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, StreamingResponse
from langchain_core.messages import AIMessageChunk, HumanMessage, ToolMessage
from prometheus_client import CONTENT_TYPE_LATEST, generate_latest
from psycopg.rows import dict_row

from .agent import AssistantAgent
from .attachments import build_attachment_blocks, get_attachment, save_attachment
from .config import AssistantConfig
from .conversations import (
    auto_title,
    create_conversation,
    delete_conversation,
    get_conversation,
    increment_message_count,
    list_conversations,
    restore_conversation,
    update_application_context,
    update_conversation,
)
from .models import (
    ChatRequest,
    ConversationCreate,
    ConversationUpdate,
)

_log = logging.getLogger("pattadar.assistant.main")


# Matches all tool/function XML blocks the model may embed in text
_XML_BLOCK_RE = re.compile(
    r"<(?:tool_call|tool_response|function_calls?|function_results?|function_response|invoke)\b[^>]*>.*?"
    r"</\s*(?:tool_call|tool_response|function_calls?|function_results?|function_response|invoke)>",
    re.DOTALL,
)


def _strip_xml_blocks(text: str) -> str:
    """Remove <tool_call> and <function_calls> XML blocks from text."""
    return _XML_BLOCK_RE.sub("", text).strip()


def _extract_text(content) -> str:
    """Extract plain text from LangChain message content (handles Anthropic list format).

    Strips raw <tool_call> and <function_calls> blocks that models emit as text.
    """
    if isinstance(content, str):
        return _strip_xml_blocks(content)
    if isinstance(content, list):
        text = "".join(
            block.get("text", "")
            for block in content
            if isinstance(block, dict) and block.get("type") == "text"
        )
        return _strip_xml_blocks(text)
    return _strip_xml_blocks(str(content))


def _extract_office_text(path: str, mime: str, fname: str) -> str | None:
    """Best-effort text extraction for office docs Claude can't read natively
    (docx/xlsx). Returns text, or None if unsupported / extraction failed."""
    low = fname.lower()
    try:
        if low.endswith(".docx") or "wordprocessingml" in mime:
            import docx  # python-docx

            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts).strip() or None
        if low.endswith(".xlsx") or "spreadsheetml" in mime:
            import openpyxl

            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            out: list[str] = []
            for ws in wb.worksheets:
                out.append(f"# Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    vals = [str(c) for c in row if c is not None]
                    if vals:
                        out.append(" | ".join(vals))
            return "\n".join(out).strip() or None
    except Exception as e:  # noqa: BLE001
        _log.warning("Office-doc extract failed for %s: %s", fname, e)
    return None


config = AssistantConfig.from_env()
agent_manager = AssistantAgent(config)


async def _get_conn() -> psycopg.AsyncConnection:
    """Get an async database connection."""
    conn = await psycopg.AsyncConnection.connect(
        config.db_uri, autocommit=True, row_factory=dict_row,
    )
    return conn


# Idempotent DDL — lets the service boot against a fresh database. The
# LangGraph checkpoint tables are managed by AsyncPostgresSaver.setup().
_DDL = """
CREATE TABLE IF NOT EXISTS r_conversations (
    id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    user_id         TEXT NOT NULL,
    title           TEXT,
    model           TEXT DEFAULT 'claude-sonnet-4-6',
    status          TEXT DEFAULT 'active' CHECK (status IN ('active','archived','deleted')),
    app_context     JSONB DEFAULT '{}',
    message_count   INT DEFAULT 0,
    created_at      TIMESTAMPTZ DEFAULT now(),
    updated_at      TIMESTAMPTZ DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_rconv_user ON r_conversations(user_id, updated_at DESC);

CREATE TABLE IF NOT EXISTS r_attachments (
    id              UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    conversation_id UUID NOT NULL REFERENCES r_conversations(id) ON DELETE CASCADE,
    user_id         TEXT NOT NULL,
    file_name       TEXT NOT NULL,
    mime_type       TEXT,
    file_size       INT,
    storage_path    TEXT NOT NULL,
    attachment_type TEXT DEFAULT 'file' CHECK (attachment_type IN ('screenshot','file','image')),
    metadata        JSONB DEFAULT '{}',
    created_at      TIMESTAMPTZ DEFAULT now()
);
CREATE INDEX IF NOT EXISTS idx_rattach_conv ON r_attachments(conversation_id);
"""


async def _ensure_tables() -> None:
    conn = await _get_conn()
    try:
        await conn.execute(_DDL)
    finally:
        await conn.close()


@asynccontextmanager
async def lifespan(app: FastAPI):
    _log.info("Assistant service starting on port %d", config.port)
    await _ensure_tables()
    # Live model registry — reads the platform_models catalog from PG with
    # a 30s refresh; falls back to a direct Anthropic fetch, then a small
    # hardcoded list.
    from . import model_registry
    registry = model_registry.init_registry(config.anthropic_api_key)
    await registry.start()
    await agent_manager.initialize()
    yield
    await agent_manager.shutdown()
    await registry.stop()
    _log.info("Assistant service shut down")


app = FastAPI(
    title="Pattadar Assistant",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

from . import telemetry as _metrics


def _user_id(x_user_id: str = Header(default="", alias="x-user-id")) -> str:
    uid = x_user_id.strip()
    if not uid:
        raise HTTPException(401, "Missing x-user-id header")
    return uid


# ── Health ─────────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    errors = []

    # 1. Check database connectivity
    try:
        conn = await _get_conn()
        try:
            await conn.execute("SELECT 1")
        finally:
            await conn.close()
    except Exception as e:
        errors.append(f"db: {e}")

    # 2. Check checkpointer is initialized and its tables exist
    if agent_manager._checkpointer is None:
        errors.append("checkpointer: not initialized")
    else:
        try:
            async with agent_manager._checkpointer._cursor() as cur:
                await cur.execute(
                    "SELECT EXISTS(SELECT 1 FROM information_schema.tables WHERE table_name = 'checkpoints')"
                )
                row = await cur.fetchone()
                if not row or not list(row.values())[0]:
                    errors.append("checkpointer: 'checkpoints' table missing — run setup()")
        except Exception as e:
            errors.append(f"checkpointer: {e}")

    if errors:
        return {"status": "unhealthy", "service": "assistant", "errors": errors}
    return {"status": "ok", "service": "assistant"}


@app.get("/metrics")
async def metrics():
    return Response(generate_latest(), media_type=CONTENT_TYPE_LATEST)


@app.get("/api/mcp-servers")
async def list_mcp_servers():
    """Return available MCP servers derived from loaded tools, plus system
    status metadata (model, available models, MCP connectivity).

    With MCP_URL unset (the v1 default) the groups list is empty and the
    status still reports the model catalog — the frontend can use this for
    a model picker later.
    """
    await agent_manager._ensure_mcp_tools()

    # Derive namespaces from actual loaded tools
    namespaces: set[str] = set()
    for tool in agent_manager._mcp_tools:
        parts = tool.name.split("_", 1)
        if len(parts) == 2:
            namespaces.add(parts[0])

    total_tools = 0
    servers = []
    for ns in sorted(namespaces):
        tool_count = sum(1 for t in agent_manager._mcp_tools if t.name.startswith(f"{ns}_"))
        total_tools += tool_count
        servers.append(
            {"id": ns, "name": ns.replace("_", " ").title(), "description": f"{ns} MCP server", "tools": tool_count}
        )
    groups = [{"label": "Tools", "servers": servers}] if servers else []

    # System status metadata
    from . import model_registry
    available_models = [
        {"id": m.id, "name": m.name, "tier": m.tier, "family": m.family, "full_name": m.full_name}
        for m in model_registry.get_registry().list_models()
    ]
    status = {
        "model": agent_manager.config.model,
        "available_models": available_models,
        "mcp": {
            "connected": len(agent_manager._mcp_tools) > 0,
            "server_count": len(namespaces),
            "tool_count": total_tools,
            "url": agent_manager.config.mcp_url or "",
        },
    }

    return {"groups": groups, "status": status}


@app.get("/docs-info")
async def docs_info():
    return {
        "name": "Pattadar Assistant",
        "version": "1.0.0",
        "description": "In-app AI assistant for Pattadar — land and property records in plain language",
        "endpoints": {
            "chat": "POST /api/chat/stream",
            "conversations": "GET /api/conversations",
            "health": "GET /health",
        },
    }


# ── Conversations CRUD ────────────────────────────────────────────────────────

@app.get("/api/conversations")
async def list_conversations_endpoint(
    x_user_id: str = Header(default="", alias="x-user-id"),
    status: str = "active",
    limit: int = 50,
    offset: int = 0,
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        rows = await list_conversations(conn, user_id, status, limit, offset)
        return {"conversations": rows}
    finally:
        await conn.close()


@app.post("/api/conversations", status_code=201)
async def create_conversation_endpoint(
    body: ConversationCreate,
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        row = await create_conversation(
            conn, user_id, body.title, body.model, body.application_context,
        )
        _metrics.conversations_total.inc()
        return row
    finally:
        await conn.close()


@app.get("/api/conversations/{conversation_id}")
async def get_conversation_endpoint(
    conversation_id: UUID,
    x_user_id: str = Header(default="", alias="x-user-id"),
    include_messages: bool = True,
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        conv = await get_conversation(conn, conversation_id, user_id)
        if not conv:
            raise HTTPException(404, "Conversation not found")

        result = {**conv}

        if include_messages:
            try:
                thread_cfg = {"configurable": {"thread_id": str(conversation_id)}}
                await _repair_orphaned_tool_calls(agent_manager.agent, thread_cfg)
                state = await agent_manager.agent.aget_state(thread_cfg)
                messages = []
                for msg in (state.values.get("messages") or []):
                    role = getattr(msg, "type", "unknown")
                    # Skip tool result messages — internal plumbing
                    if role == "tool":
                        continue
                    text = _extract_text(msg.content)
                    # Skip AI messages that only contained tool calls (no user-facing text)
                    if role == "ai" and not text and hasattr(msg, "tool_calls") and msg.tool_calls:
                        continue
                    m = {"role": role, "content": text}
                    if hasattr(msg, "tool_calls") and msg.tool_calls:
                        m["tool_calls"] = [
                            {"name": tc.get("name", ""), "args": tc.get("args", {})}
                            for tc in msg.tool_calls
                        ]
                    messages.append(m)
                result["messages"] = messages
            except Exception as e:
                _log.warning("Failed to load messages for %s: %s", conversation_id, e)
                result["messages"] = []

        return result
    finally:
        await conn.close()


@app.put("/api/conversations/{conversation_id}")
async def update_conversation_endpoint(
    conversation_id: UUID,
    body: ConversationUpdate,
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        row = await update_conversation(
            conn, conversation_id, user_id, body.title, body.status,
        )
        if not row:
            raise HTTPException(404, "Conversation not found")
        return row
    finally:
        await conn.close()


@app.delete("/api/conversations/{conversation_id}")
async def delete_conversation_endpoint(
    conversation_id: UUID,
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        ok = await delete_conversation(conn, conversation_id, user_id)
        if not ok:
            raise HTTPException(404, "Conversation not found")
        return {"ok": True}
    finally:
        await conn.close()


@app.post("/api/conversations/{conversation_id}/restore")
async def restore_conversation_endpoint(
    conversation_id: UUID,
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        row = await restore_conversation(conn, conversation_id, user_id)
        if not row:
            raise HTTPException(404, "Conversation not found or not deleted")
        return row
    finally:
        await conn.close()


# ── Chat Stream (SSE) ─────────────────────────────────────────────────────────

async def _repair_orphaned_tool_calls(agent, thread_config) -> int:
    """Fix corrupted checkpoint where AIMessage has tool_calls but no ToolMessages.

    LangGraph checkpoints after each node. If the tool node fails before
    completing (tool timeout, API error, client disconnect), the checkpoint
    has an AIMessage with tool_calls but no subsequent ToolMessages.
    This makes the conversation permanently stuck. Fix by injecting
    synthetic error ToolMessages so the conversation can continue.

    Returns the number of repaired tool calls.
    """
    try:
        state = await agent.aget_state(thread_config)
        messages = state.values.get("messages") or []
        if not messages:
            return 0

        # Find orphaned tool_calls: AI messages with tool_calls that aren't
        # followed by matching ToolMessages
        tool_call_ids_seen = set()
        orphaned_tool_calls = []

        for msg in messages:
            if hasattr(msg, "tool_calls") and msg.tool_calls:
                for tc in msg.tool_calls:
                    tc_id = tc.get("id", "")
                    if tc_id:
                        tool_call_ids_seen.add(tc_id)
                        orphaned_tool_calls.append(tc)
            if getattr(msg, "type", "") == "tool":
                tc_id = getattr(msg, "tool_call_id", "")
                if tc_id:
                    tool_call_ids_seen.discard(tc_id)
                    orphaned_tool_calls = [
                        tc for tc in orphaned_tool_calls if tc.get("id") != tc_id
                    ]

        if not orphaned_tool_calls:
            return 0

        repair_messages = [
            ToolMessage(
                content=f"Error: Tool '{tc.get('name', 'unknown')}' failed to execute. The operation did not complete. Please retry if needed.",
                tool_call_id=tc.get("id", ""),
                name=tc.get("name", ""),
            )
            for tc in orphaned_tool_calls
        ]
        await agent.aupdate_state(thread_config, {"messages": repair_messages})
        _log.warning(
            "Repaired %d orphaned tool_call(s) in thread %s: %s",
            len(repair_messages),
            thread_config["configurable"]["thread_id"],
            [tc.get("name") for tc in orphaned_tool_calls],
        )
        return len(repair_messages)
    except Exception as e:
        _log.warning(
            "Failed to repair state for %s: %s",
            thread_config["configurable"]["thread_id"], e,
        )
        return 0


@app.post("/api/chat/stream")
async def chat_stream(
    body: ChatRequest,
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)

    # Verify conversation ownership
    try:
        conn = await _get_conn()
    except Exception as e:
        _log.exception("DB connection failed for chat_stream")
        raise HTTPException(503, f"Database unavailable: {e}")
    try:
        conv = await get_conversation(conn, body.conversation_id, user_id)
        if not conv:
            raise HTTPException(404, "Conversation not found")

        # Auto-title from first message
        await auto_title(conn, body.conversation_id, user_id, body.message)
        await increment_message_count(conn, body.conversation_id)

        # Persist latest application context snapshot on the conversation
        if body.application_context and (
            body.application_context.get("app") or body.application_context.get("insights")
            or body.application_context.get("components") or body.application_context.get("navigation")
        ):
            await update_application_context(
                conn, str(body.conversation_id), user_id, body.application_context
            )
    except HTTPException:
        raise
    except Exception as e:
        _log.exception("Conversation setup error for %s", body.conversation_id)
        raise HTTPException(500, f"Failed to prepare conversation: {e}")
    finally:
        await conn.close()

    # Build message content — plain text or multimodal with attachments
    content: str | list[dict] = body.message
    if body.attachment_ids:
        att_conn = await _get_conn()
        try:
            content_blocks, _image_attachments = await build_attachment_blocks(
                att_conn, body.attachment_ids, user_id, _extract_office_text
            )
            if content_blocks:
                content_blocks.append({"type": "text", "text": body.message})
                content = content_blocks
        finally:
            await att_conn.close()

    # Hot-reload prompt from DB before building the agent
    await agent_manager.refresh_prompt()

    # Build contextual prompt with live application context
    contextual_prompt = agent_manager._build_contextual_prompt(body.application_context)
    _log.info(
        "chat_stream context: app=%s, insights=%d, nav=%d, prompt_len=%d",
        (body.application_context.get("app") or {}).get("name", "none"),
        len(body.application_context.get("insights") or body.application_context.get("components", [])),
        len(body.application_context.get("navigation", [])),
        len(contextual_prompt),
    )

    model = body.model or conv.get("model") or None
    navigation = body.application_context.get("navigation", []) if body.application_context else []
    await agent_manager._ensure_mcp_tools()
    agent = agent_manager.get_agent_with_servers(model, body.enabled_servers, navigation=navigation or None)

    # Pass contextual prompt via config so the dynamic prompt function picks it up
    thread_config = {
        "configurable": {
            "thread_id": str(body.conversation_id),
            "system_prompt": contextual_prompt,
        },
    }

    # Repair any corrupted checkpoint state before streaming
    await _repair_orphaned_tool_calls(agent, thread_config)

    # Record chat metrics
    _used_model = model or "default"
    _metrics.chat_total.labels(model=_used_model).inc()
    import time as _time
    _chat_t0 = _time.perf_counter()

    async def generate() -> AsyncGenerator[str, None]:
        yield ": keepalive\n\n"

        # Stateful filter: suppress XML tool blocks the model emits as text
        # when tools aren't bound (<function_calls>, <function_response>,
        # <tool_call>, <invoke>, etc.).
        _xml_buf = ""
        _suppressing = False
        _end_tag = ""
        _OPEN_RE = re.compile(r"<(function_calls?|function_results?|function_response|tool_call|tool_response|invoke)\b")
        # Pre-compute all possible prefixes of suppressible tags for partial matching.
        _TAG_STRS = [
            "<function_calls>", "<function_call>", "<function_results>",
            "<function_result>", "<function_response>",
            "<tool_call>", "<tool_response>", "<invoke ",
        ]
        _PREFIXES: set[str] = set()
        for _t in _TAG_STRS:
            for _i in range(1, len(_t) + 1):
                _PREFIXES.add(_t[:_i])
        _MAX_PREFIX = max(len(p) for p in _PREFIXES)

        def _filter_text(text: str) -> str:
            """Feed text through the XML block filter. Returns safe text to emit."""
            nonlocal _xml_buf, _suppressing, _end_tag

            _xml_buf += text
            parts: list[str] = []

            while _xml_buf:
                if _suppressing:
                    end_idx = _xml_buf.find(_end_tag)
                    if end_idx >= 0:
                        _suppressing = False
                        _xml_buf = _xml_buf[end_idx + len(_end_tag):]
                        continue
                    # Check if tail could be partial end tag
                    for i in range(min(len(_end_tag) - 1, len(_xml_buf)), 0, -1):
                        if _xml_buf.endswith(_end_tag[:i]):
                            _xml_buf = _xml_buf[-i:]
                            return "".join(parts)
                    _xml_buf = ""
                    return "".join(parts)

                # Not suppressing — look for opening tag
                m = _OPEN_RE.search(_xml_buf)
                if m:
                    parts.append(_xml_buf[:m.start()])
                    tag_name = m.group(1)
                    _end_tag = f"</{tag_name}>"
                    _suppressing = True
                    # Skip past the closing ">" of the opening tag
                    rest = _xml_buf[m.end():]
                    gt = rest.find(">")
                    _xml_buf = rest[gt + 1:] if gt >= 0 else ""
                    continue

                # No full match — check if tail is a partial prefix of a tag
                found_partial = False
                for i in range(min(_MAX_PREFIX, len(_xml_buf)), 0, -1):
                    if _xml_buf[-i:] in _PREFIXES:
                        parts.append(_xml_buf[:-i])
                        _xml_buf = _xml_buf[-i:]
                        found_partial = True
                        break

                if not found_partial:
                    parts.append(_xml_buf)
                    _xml_buf = ""
                break  # exit loop after non-suppressing path

            return "".join(parts)

        try:
            # Use astream_events (not astream) to get raw model-level token
            # events. LangGraph's astream(stream_mode="messages") buffers
            # thinking blocks into one large chunk; astream_events gives us
            # each thinking_delta individually for progressive streaming.
            async for event in agent.astream_events(
                {"messages": [HumanMessage(content=content)]},
                thread_config,
                version="v2",
            ):
                kind = event["event"]

                # Raw model token — each AIMessageChunk has one thinking
                # delta OR one text delta, giving true progressive streaming
                if kind == "on_chat_model_stream":
                    chunk = event["data"]["chunk"]
                    if isinstance(chunk, AIMessageChunk) and chunk.content:
                        if isinstance(chunk.content, str):
                            safe = _filter_text(chunk.content)
                            if safe:
                                yield f"data: {json.dumps({'type': 'token', 'text': safe})}\n\n"
                        elif isinstance(chunk.content, list):
                            for block in chunk.content:
                                if not isinstance(block, dict):
                                    continue
                                if block.get("type") == "thinking" and block.get("thinking"):
                                    yield f"data: {json.dumps({'type': 'thinking', 'text': block['thinking']})}\n\n"
                                elif block.get("type") == "text" and block.get("text"):
                                    safe = _filter_text(block["text"])
                                    if safe:
                                        yield f"data: {json.dumps({'type': 'token', 'text': safe})}\n\n"

                    # Tool call chunks from model (structured tool_use)
                    if isinstance(chunk, AIMessageChunk) and chunk.tool_call_chunks:
                        for tc in chunk.tool_call_chunks:
                            if tc.get("name"):
                                yield f"data: {json.dumps({'type': 'tool_start', 'name': tc['name'], 'id': tc.get('id', '')})}\n\n"

                # Tool finished executing
                elif kind == "on_tool_end":
                    tool_output = event["data"].get("output", "")
                    tool_name = event.get("name", "")
                    # Record tool call metric — MCP tools are "{namespace}_{tool}";
                    # built-in tools (navigate_user etc.) land under their first segment.
                    _idx = tool_name.find("_")
                    _srv = tool_name[:_idx] if _idx > 0 else "unknown"
                    _tl = tool_name[_idx + 1:] if _idx > 0 else tool_name
                    _metrics.tool_calls_total.labels(server=_srv, tool=_tl).inc()
                    yield f"data: {json.dumps({'type': 'tool_end', 'name': tool_name, 'content': str(tool_output)[:500]})}\n\n"

                    # Emit an action event for any tool that returns an
                    # {"action", ...} envelope (navigate_user + the page-action
                    # tools set_filter / open_record / fill_field / submit_form).
                    # The browser forwards these to the open app.
                    try:
                        action_result = tool_output if isinstance(tool_output, dict) else {}
                        if not action_result:
                            raw = getattr(tool_output, "content", str(tool_output))
                            if isinstance(raw, str) and "{" in raw:
                                action_result = json.loads(raw)
                        if isinstance(action_result, dict) and action_result.get("action"):
                            action_name = action_result["action"]
                            if action_name == "navigate":
                                # navigate carries `path` at top level.
                                if action_result.get("path"):
                                    _log.info("Emitting navigate action SSE: path=%s", action_result["path"])
                                    yield f"data: {json.dumps({'type': 'action', 'action': 'navigate', 'path': action_result['path']})}\n\n"
                            else:
                                _log.info("Emitting page action SSE: action=%s", action_name)
                                yield f"data: {json.dumps({'type': 'action', 'action': action_name, 'args': action_result.get('args', {})})}\n\n"
                    except Exception as exc:
                        _log.warning("action parse error for %s: %s", tool_name, exc)

            _metrics.chat_seconds.labels(model=_used_model).observe(_time.perf_counter() - _chat_t0)
            yield "data: [DONE]\n\n"
        except Exception as e:
            _log.exception("Stream error for conversation %s", body.conversation_id)
            _metrics.chat_errors_total.labels(error_type=type(e).__name__).inc()
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ── Attachments ────────────────────────────────────────────────────────────────

@app.post("/api/attachments", status_code=201)
async def upload_attachment(
    file: UploadFile = File(...),
    conversation_id: str = Form(...),
    attachment_type: str = Form("file"),
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)
    file_bytes = await file.read()

    if len(file_bytes) > config.max_upload_bytes:
        raise HTTPException(413, f"File too large (max {config.max_upload_bytes // 1024 // 1024} MB)")

    conn = await _get_conn()
    try:
        conv = await get_conversation(conn, UUID(conversation_id), user_id)
        if not conv:
            raise HTTPException(404, "Conversation not found")

        row = await save_attachment(
            conn, conversation_id, user_id,
            file.filename or "upload",
            file_bytes, file.content_type, attachment_type,
            config.upload_dir,
        )
        return row
    finally:
        await conn.close()


@app.get("/api/attachments/{attachment_id}")
async def download_attachment(
    attachment_id: UUID,
    x_user_id: str = Header(default="", alias="x-user-id"),
):
    user_id = _user_id(x_user_id)
    conn = await _get_conn()
    try:
        row = await get_attachment(conn, str(attachment_id), user_id)
        if not row:
            raise HTTPException(404, "Attachment not found")
        return FileResponse(
            row["storage_path"],
            filename=row["file_name"],
            media_type=row.get("mime_type", "application/octet-stream"),
        )
    finally:
        await conn.close()
